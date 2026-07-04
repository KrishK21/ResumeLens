"""
ResumeLens — FastAPI application
================================

Three endpoints mirroring the user flow:

  POST /upload   multipart .docx  -> session_id + located bullets
  POST /tailor   {session_id, job_description} -> keywords, gaps, rewrites
  POST /export   {session_id, default_tone, choices, format} -> file download

Run locally (from the backend folder, venv active):
    uvicorn app.main:app --reload --port 8000
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from docx import Document

from .config import get_settings
from .schemas import (
    UploadResponse, BulletOut,
    TailorRequest, TailorResponse, KeywordMatchOut, RewriteOut,
    ExportRequest,
)
from .document_engine import (
    find_experience_bullets, apply_rewrites, docx_to_pdf,
    assert_text_is_selectable,
)
from .llm import (
    extract_keywords, match_keywords_against_resume, rewrite_bullet,
    shorten_bullet, shorten_many,
)
from .fit_engine import fit_to_one_page, rank_bullets_by_weakness
from .store import store
from . import profile_store

settings = get_settings()

app = FastAPI(title="ResumeLens API", version="0.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.cors_origin_list,   # only the configured frontend
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# --------------------------------------------------------------------------- #
#  Health
# --------------------------------------------------------------------------- #

@app.get("/health")
def health() -> dict:
    return {"status": "ok"}


# --------------------------------------------------------------------------- #
#  /upload
# --------------------------------------------------------------------------- #

@app.post("/upload", response_model=UploadResponse)
async def upload(file: UploadFile = File(...)) -> UploadResponse:
    data = await _read_upload(file)
    name = file.filename or "resume.docx"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)
    try:
        return _session_from_docx(tmp_path, name)
    finally:
        tmp_path.unlink(missing_ok=True)


async def _read_upload(file: UploadFile) -> bytes:
    """Validate + read an uploaded .docx with a hard size cap."""
    name = file.filename or "resume.docx"
    if not name.lower().endswith(".docx"):
        raise HTTPException(400, "Only .docx files are supported right now.")
    max_bytes = settings.max_upload_mb * 1024 * 1024
    data = await file.read(max_bytes + 1)
    if len(data) > max_bytes:
        raise HTTPException(413, f"File exceeds {settings.max_upload_mb} MB limit.")
    if not data:
        raise HTTPException(400, "Empty file.")
    return data


def _session_from_docx(docx_path: Path, name: str) -> UploadResponse:
    """Parse a .docx into a session + located bullets (shared by upload & profile)."""
    try:
        doc = Document(str(docx_path))
    except Exception:
        raise HTTPException(400, "Could not read this file as a Word document.")

    bullets = find_experience_bullets(doc)
    unsupported = False
    note = ""
    if not bullets:
        unsupported = _likely_unsupported_layout(doc)
        note = (
            "No editable experience bullets were found. If your resume uses text "
            "boxes or a multi-column template, those aren't editable yet — try a "
            "standard single-column Word resume."
            if unsupported else
            "No bulleted experience section was detected."
        )

    sess = store.create(docx_path, name, bullets)
    return UploadResponse(
        session_id=sess.id,
        filename=name,
        bullets=[BulletOut(index=b.index, text=b.text) for b in bullets],
        unsupported_layout=unsupported,
        note=note,
    )


# --------------------------------------------------------------------------- #
#  Profile — saved resume (no auth yet; keyed by a browser profile id)
# --------------------------------------------------------------------------- #

@app.post("/profile/{profile_id}/resume", response_model=dict)
async def save_profile_resume(profile_id: str, file: UploadFile = File(...)) -> dict:
    """Save (or replace) the resume stored on a profile."""
    data = await _read_upload(file)
    name = file.filename or "resume.docx"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)
    try:
        # Validate it's a readable docx before saving.
        _session_from_docx(tmp_path, name)  # raises on bad file
        profile_store.save_resume(profile_id, tmp_path, name)
    except ValueError:
        raise HTTPException(400, "Invalid profile id.")
    finally:
        tmp_path.unlink(missing_ok=True)
    return {"saved": True, "filename": name}


@app.get("/profile/{profile_id}/resume", response_model=dict)
def get_profile_resume(profile_id: str) -> dict:
    """Check whether the profile has a saved resume."""
    path = profile_store.get_resume_path(profile_id)
    if path is None:
        return {"has_resume": False, "filename": None}
    return {"has_resume": True, "filename": profile_store.get_resume_filename(profile_id)}


@app.delete("/profile/{profile_id}/resume", response_model=dict)
def delete_profile_resume(profile_id: str) -> dict:
    profile_store.delete_resume(profile_id)
    return {"deleted": True}


@app.post("/profile/{profile_id}/session", response_model=UploadResponse)
def session_from_profile(profile_id: str) -> UploadResponse:
    """Start a tailoring session from the profile's SAVED resume (no re-upload)."""
    path = profile_store.get_resume_path(profile_id)
    if path is None:
        raise HTTPException(404, "No saved resume on this profile.")
    name = profile_store.get_resume_filename(profile_id) or "resume.docx"
    return _session_from_docx(path, name)


def _likely_unsupported_layout(doc: Document) -> bool:
    """Heuristic: body has almost no paragraph text but the file clearly has
    content (e.g. in tables) -> probably a text-box/table template."""
    body_text = "".join(p.text for p in doc.paragraphs).strip()
    has_tables = len(doc.tables) > 0
    return len(body_text) < 50 or has_tables


# --------------------------------------------------------------------------- #
#  /tailor   (extract keywords + rewrite all bullets)
# --------------------------------------------------------------------------- #

@app.post("/tailor", response_model=TailorResponse)
def tailor(req: TailorRequest) -> TailorResponse:
    sess = store.get(req.session_id)
    if sess is None:
        raise HTTPException(404, "Session not found or expired. Re-upload your resume.")
    if not sess.bullets:
        raise HTTPException(400, "This resume has no editable bullets to tailor.")

    # Stage 1: keywords + gap analysis
    kw = extract_keywords(req.job_description)
    doc = Document(str(sess.docx_path))
    resume_text = "\n".join(p.text for p in doc.paragraphs)
    matches = match_keywords_against_resume(kw, resume_text)
    gaps = [m.term for m in matches if not m.matched]

    # Stage 2: rewrite every bullet (eager) -> store for /export
    target_terms = kw.all_terms()
    rewrites_out: list[RewriteOut] = []
    rewrites_store: dict[int, object] = {}
    for b in sess.bullets:
        rw = rewrite_bullet(b.text, target_terms)
        rewrites_store[b.index] = rw
        rewrites_out.append(RewriteOut(
            index=b.index,
            original=rw.original,
            conservative=rw.conservative,
            balanced=rw.balanced,
            keywords_used_conservative=rw.keywords_used_conservative,
            keywords_used_balanced=rw.keywords_used_balanced,
            fabrication_check=rw.fabrication_check,
        ))
    store.set_rewrites(sess.id, rewrites_store)

    return TailorResponse(
        job_title=kw.job_title,
        keywords=[KeywordMatchOut(term=m.term, matched=m.matched) for m in matches],
        gaps=gaps,
        rewrites=rewrites_out,
    )


# --------------------------------------------------------------------------- #
#  /export   (apply chosen tones -> downloadable file)
# --------------------------------------------------------------------------- #

@app.post("/export")
def export(req: ExportRequest):
    sess = store.get(req.session_id)
    if sess is None:
        raise HTTPException(404, "Session not found or expired. Re-upload your resume.")
    if not sess.rewrites:
        raise HTTPException(400, "Nothing to export yet — run /tailor first.")

    # Build the final text for each bullet from the user's choices.
    # Any bullet without an explicit choice uses default_tone (auto mode).
    # Choices referencing unknown indices are ignored (defensive).
    valid_indices = set(sess.rewrites.keys())
    overrides = {c.index: c for c in req.choices if c.index in valid_indices}
    final: dict[int, str] = {}
    for index, rw in sess.rewrites.items():
        choice = overrides.get(index)
        if choice is None:
            final[index] = rw.version(req.default_tone)
        elif choice.tone == "custom" and choice.custom_text is not None:
            final[index] = choice.custom_text
        else:
            final[index] = rw.version(choice.tone)

    out_dir = sess.docx_path.parent / "out"
    try:
        # --- One-page fit (text-only; formatting never changed) ---
        fit_headers: dict[str, str] = {}
        if req.one_page:
            # Compact aims for a tighter target so the fit loop shortens more,
            # but we no longer pre-shrink every bullet upfront (that was slow).
            # The fit loop only shortens bullets that actually cause overflow,
            # now in parallel batches for speed.
            rounds = 4 if req.fit_mode == "compact" else 3
            weakness = rank_bullets_by_weakness(final)
            final, fit_report = fit_to_one_page(
                sess.docx_path, final, out_dir,
                shorten_many_fn=shorten_many,
                weakness_rank=weakness,
                max_shorten_rounds=rounds,
                allow_drop=True,
            )
            fit_headers = {
                "X-Fit-Pages": str(fit_report.final_pages),
                "X-Fit-OnePage": str(fit_report.fits_one_page).lower(),
                "X-Fit-Shortened": ",".join(map(str, fit_report.shortened_indices)),
                "X-Fit-Dropped": ",".join(map(str, fit_report.dropped_indices)),
            }

        edited_docx = apply_rewrites(sess.docx_path, final, out_dir / "tailored.docx")

        if req.format == "docx":
            return FileResponse(
                str(edited_docx),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename="resume_tailored.docx",
                headers={
                    **fit_headers,
                    # Force a download, never an inline preview.
                    "Content-Disposition": 'attachment; filename="resume_tailored.docx"',
                },
            )

        # PDF path: export AND verify the text layer is selectable before returning.
        pdf = docx_to_pdf(edited_docx, out_dir)
        assert_text_is_selectable(pdf)
        return FileResponse(
            str(pdf),
            media_type="application/pdf",
            filename="resume_tailored.pdf",
            headers={
                **fit_headers,
                # Force a download, never an inline browser preview (which on
                # some Chrome/Windows setups can auto-open a print prompt).
                "Content-Disposition": 'attachment; filename="resume_tailored.pdf"',
            },
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Export failed: {type(e).__name__}: {e}")